"""
Обработка документов для RAG: извлечение, модерация, нормализация и чанкинг.
Поддержка: txt, doc, docx, xls, xlsx, pdf, html.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict, Any
import hashlib
import mimetypes
import os
import re
import unicodedata

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {".txt", ".doc", ".docx", ".xls", ".xlsx", ".pdf", ".html", ".htm"}


@dataclass
class ProcessedDocument:
    source_path: str
    source_name: str
    extension: str
    file_hash: str
    normalized_text: str
    chunks: List[str]


class DocumentProcessor:
    """Единый пайплайн обработки пользовательских документов."""

    def __init__(self, max_file_size_mb: int = 25, enable_ocr: bool = True):
        self.max_file_size_mb = max_file_size_mb
        self.enable_ocr = enable_ocr

    def validate_file(self, file_path: str) -> None:
        path = Path(file_path)
        if not path.exists() or not path.is_file():
            raise FileNotFoundError(f"Файл не найден: {file_path}")

        ext = path.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Неподдерживаемый формат: {ext}. Допустимо: {sorted(SUPPORTED_EXTENSIONS)}")

        max_bytes = self.max_file_size_mb * 1024 * 1024
        file_size = path.stat().st_size
        if file_size > max_bytes:
            raise ValueError(
                f"Файл {path.name} слишком большой: {file_size / (1024 * 1024):.2f} MB. "
                f"Лимит: {self.max_file_size_mb} MB"
            )

        guessed_type, _ = mimetypes.guess_type(str(path))
        if guessed_type and not any(
            token in guessed_type for token in ["text", "word", "excel", "sheet", "pdf", "html", "xml"]
        ):
            raise ValueError(f"Подозрительный MIME тип: {guessed_type}")

    def process_file(self, file_path: str, chunk_size: int = 700, overlap: int = 120) -> ProcessedDocument:
        self.validate_file(file_path)

        path = Path(file_path)
        ext = path.suffix.lower()

        raw_text = self._extract_text(path, ext)
        normalized = self._normalize_text(raw_text)
        chunks = self._chunk_text(normalized, chunk_size=chunk_size, overlap=overlap)

        if not chunks:
            raise ValueError(f"После обработки файл {path.name} не содержит пригодного текста")

        file_hash = self._compute_file_hash(path)

        return ProcessedDocument(
            source_path=str(path.resolve()),
            source_name=path.name,
            extension=ext,
            file_hash=file_hash,
            normalized_text=normalized,
            chunks=chunks,
        )

    def _extract_text(self, path: Path, ext: str) -> str:
        if ext == ".txt":
            return path.read_text(encoding="utf-8", errors="ignore")

        if ext in {".html", ".htm"}:
            html = path.read_text(encoding="utf-8", errors="ignore")
            soup = BeautifulSoup(html, "html.parser")
            return soup.get_text("\n")

        if ext == ".docx":
            doc = DocxDocument(str(path))
            return "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())

        if ext == ".doc":
            return self._extract_doc_legacy(path)

        if ext in {".xls", ".xlsx"}:
            return self._extract_excel(path)

        if ext == ".pdf":
            return self._extract_pdf(path)

        raise ValueError(f"Неподдерживаемое расширение: {ext}")

    def _extract_doc_legacy(self, path: Path) -> str:
        # Для .doc используем textract как кросс-платформенный вариант.
        try:
            import textract  # type: ignore
        except Exception as exc:
            raise RuntimeError(
                "Для чтения .doc нужен пакет textract. Установите зависимости и системные утилиты "
                "(например antiword) или конвертируйте файл в .docx."
            ) from exc

        data = textract.process(str(path))
        return data.decode("utf-8", errors="ignore")

    def _extract_excel(self, path: Path) -> str:
        sheets = pd.read_excel(path, sheet_name=None)
        blocks: List[str] = []

        for sheet_name, df in sheets.items():
            if df is None or df.empty:
                continue

            cleaned_df = df.fillna("")
            rows = []
            for row in cleaned_df.astype(str).values.tolist():
                row_text = " | ".join(cell.strip() for cell in row if str(cell).strip())
                if row_text:
                    rows.append(row_text)

            if rows:
                blocks.append(f"[Лист: {sheet_name}]\n" + "\n".join(rows))

        return "\n\n".join(blocks)

    def _extract_pdf(self, path: Path) -> str:
        reader = PdfReader(str(path))
        pages: List[str] = []

        for page in reader.pages:
            page_text = page.extract_text() or ""
            pages.append(page_text)

        text = "\n".join(pages).strip()

        # OCR fallback для сканов (локально, без внешних облаков).
        if len(text) < 100 and self.enable_ocr:
            ocr_text = self._ocr_pdf(path)
            if ocr_text.strip():
                return ocr_text

        return text

    def _ocr_pdf(self, path: Path) -> str:
        try:
            from pdf2image import convert_from_path
            import pytesseract
        except Exception as exc:
            raise RuntimeError(
                "OCR для PDF не готов: нужны пакеты pdf2image и pytesseract, "
                "а также установленный Tesseract OCR."
            ) from exc

        lang = os.getenv("OCR_LANG", "rus+eng")
        images = convert_from_path(str(path), dpi=250)

        text_parts: List[str] = []
        for img in images:
            page_text = pytesseract.image_to_string(img, lang=lang)
            if page_text:
                text_parts.append(page_text)

        return "\n".join(text_parts)

    def _normalize_text(self, text: str) -> str:
        text = unicodedata.normalize("NFKC", text)
        text = text.replace("\xa0", " ")

        # Удаляем артефакты управляющих символов, сохраняем переносы и табы.
        text = "".join(ch for ch in text if ch == "\n" or ch == "\t" or ch.isprintable())

        # Бережная нормализация для RU-EN: не переводим регистр и не трогаем бренды.
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _chunk_text(self, text: str, chunk_size: int = 700, overlap: int = 120) -> List[str]:
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        chunks: List[str] = []
        current = ""

        for paragraph in paragraphs:
            if len(current) + len(paragraph) + 2 <= chunk_size:
                current = f"{current}\n\n{paragraph}".strip() if current else paragraph
                continue

            if current:
                chunks.append(current)
                overlap_text = self._build_overlap_tail(current, overlap)
                current = f"{overlap_text}\n\n{paragraph}".strip()
            else:
                sentence_chunks = self._split_long_paragraph(paragraph, chunk_size, overlap)
                if sentence_chunks:
                    chunks.extend(sentence_chunks[:-1])
                    current = sentence_chunks[-1]

        if current:
            chunks.append(current)

        return [chunk for chunk in chunks if len(chunk) >= 50]

    def _split_long_paragraph(self, paragraph: str, chunk_size: int, overlap: int) -> List[str]:
        sentences = re.split(r"(?<=[.!?])\s+", paragraph)

        chunks: List[str] = []
        current = ""

        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue

            if len(current) + len(sentence) + 1 <= chunk_size:
                current = f"{current} {sentence}".strip() if current else sentence
                continue

            if current:
                chunks.append(current)
                overlap_text = self._build_overlap_tail(current, overlap)
                current = f"{overlap_text} {sentence}".strip()
            else:
                current = sentence

        if current:
            chunks.append(current)

        return chunks

    @staticmethod
    def _build_overlap_tail(text: str, overlap: int) -> str:
        """Формирует overlap по границам предложений/слов, чтобы не резать лексемы."""
        if not text:
            return ""
        if overlap <= 0:
            return ""
        if len(text) <= overlap:
            return text.strip()

        tail = text[-overlap:]

        # Приоритет: начать overlap после последнего завершенного предложения.
        sentence_breaks = list(re.finditer(r"[.!?]\s+", tail))
        if sentence_breaks:
            candidate = tail[sentence_breaks[-1].end():].strip()
            if candidate:
                return candidate

        # Фолбэк: начать после ближайшего пробела (граница слова).
        space_idx = tail.find(" ")
        if space_idx != -1 and space_idx + 1 < len(tail):
            candidate = tail[space_idx + 1 :].strip()
            if candidate:
                return candidate

        return tail.strip()

    @staticmethod
    def _compute_file_hash(path: Path) -> str:
        hasher = hashlib.sha256()
        with path.open("rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                hasher.update(chunk)
        return hasher.hexdigest()
